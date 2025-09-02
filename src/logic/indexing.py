"""
ChromaDBインデックス機能実装 (TDD Green フェーズ)
PDF/TXTファイル読み込み・ベクトル化・インデックス管理機能を提供
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import PyPDF2
import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings

from src.models.document import Document
from src.exceptions.base_exceptions import IndexingError
from src.utils.structured_logger import get_logger
from src.utils.progress_utils import ProgressTracker, should_show_progress
from src.utils.cancellation_utils import CancellableOperation
from src.security.file_validator import FileValidator


class ChromaDBIndexer(CancellableOperation):
    """
    ChromaDB ベクトルデータベースインデクサー
    
    PDF/TXTファイルを読み込み、ベクトル化してインデックスを作成・管理する
    """
    
    # 埋め込みモデル次元数マッピング
    EMBEDDING_DIMENSIONS = {
        "nomic-embed-text": 768,
        "mxbai-embed-large": 1024,
        "all-minilm": 384,
        "snowflake-arctic-embed": 1024
    }
    
    def __init__(
        self,
        collection_name: str = "documents",
        db_path: str = "./data/chroma_db",
        embedding_model: str = "nomic-embed-text"
    ):
        """
        ChromaDBインデクサーを初期化
        
        Args:
            collection_name: コレクション名
            db_path: データベースパス
            embedding_model: 埋め込みモデル名 (Ollama)
        """
        super().__init__(f"ChromaDB Indexer ({collection_name})")
        
        self.collection_name = collection_name
        self.db_path = Path(db_path)
        self.embedding_model = embedding_model
        self.logger = get_logger(__name__)
        
        # ChromaDB設定とエラーハンドリング強化
        try:
            self.db_path.mkdir(parents=True, exist_ok=True)
            self.client = chromadb.PersistentClient(
                path=str(self.db_path),
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
        except Exception as e:
            self.logger.error(f"ChromaDBクライアント初期化エラー: {e}", exc_info=True, extra={"db_path": str(db_path)})
            raise IndexingError(
                f"ChromaDBクライアント初期化エラー: {e}",
                error_code="IDX-000",
                details={"db_path": str(db_path), "original_error": str(e)}
            ) from e
        
        # 埋め込み関数の初期化（エラーハンドリング強化）
        self._initialize_embedding_function()
        
        # コレクション取得または作成（次元数チェック付き）
        try:
            self._initialize_collection_with_dimension_check(collection_name)
        except Exception as e:
            self.logger.error(f"ChromaDBコレクション初期化エラー: {e}", exc_info=True, 
                            extra={"collection_name": collection_name, "db_path": str(db_path)})
            raise IndexingError(
                f"ChromaDBコレクション初期化エラー: {e}",
                error_code="IDX-000",
                details={"collection_name": collection_name, "db_path": str(db_path)}
            ) from e
        
        # テキスト分割器の初期化
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # セキュリティバリデーター初期化
        self.file_validator = FileValidator()
        
        self.logger.info(f"ChromaDBIndexer初期化完了", extra={
            "collection_name": collection_name,
            "db_path": str(db_path),
            "embedding_model": embedding_model
        })
    
    def _initialize_embedding_function(self) -> None:
        """
        埋め込み関数を初期化（エラーハンドリング強化）
        
        Ollamaが利用できない場合はデフォルト埋め込みにフォールバック
        """
        try:
            # Ollama接続テスト
            import requests
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if response.status_code == 200:
                self._embedding_function = OllamaEmbeddings(
                    model=self.embedding_model,
                    base_url="http://localhost:11434"
                )
                self.logger.info(f"Ollama埋め込み初期化成功: {self.embedding_model}")
            else:
                raise ConnectionError("Ollama service not available")
                
        except Exception as e:
            self.logger.warning(
                f"Ollama埋め込み初期化失敗、デフォルト埋め込みを使用: {e}",
                extra={"embedding_model": self.embedding_model, "error": str(e)}
            )
            self._embedding_function = None
    
    def _initialize_collection_with_dimension_check(self, collection_name: str) -> None:
        """
        次元数チェック付きでコレクションを初期化
        
        Args:
            collection_name: コレクション名
        """
        try:
            # 既存コレクションの確認
            existing_collections = self.client.list_collections()
            collection_exists = any(c.name == collection_name for c in existing_collections)
            
            if collection_exists:
                # 既存コレクションを取得
                self.collection = self.client.get_collection(name=collection_name)
                
                # 次元数の互換性をチェック
                compatibility = self.check_embedding_dimension_compatibility()
                
                if not compatibility['is_compatible'] and compatibility.get('needs_recreation', False):
                    self.logger.warning(
                        f"コレクション次元数不整合を検出、再作成します: {compatibility['error']}"
                    )
                    # 不整合の場合はコレクションを削除して再作成
                    self.client.delete_collection(collection_name)
                    self._create_new_collection(collection_name)
                else:
                    self.logger.info(f"既存コレクションを使用: {collection_name}")
            else:
                # 新規コレクションを作成
                self._create_new_collection(collection_name)
                
        except Exception as e:
            self.logger.error(f"コレクション初期化エラー: {e}")
            # フォールバック: 基本的なコレクション作成を試行
            try:
                self.collection = self.client.get_or_create_collection(
                    name=collection_name,
                    metadata={"description": f"Knowledge base collection: {collection_name}"}
                )
            except Exception as fallback_error:
                raise IndexingError(
                    f"コレクション作成失敗: {fallback_error}",
                    error_code="IDX-COLLECTION-INIT",
                    details={"collection_name": collection_name, "original_error": str(e)}
                ) from fallback_error

    def _create_new_collection(self, collection_name: str) -> None:
        """
        新しいコレクションを作成
        
        Args:
            collection_name: コレクション名
        """
        expected_dimensions = self.get_model_expected_dimensions()
        
        self.collection = self.client.create_collection(
            name=collection_name,
            metadata={
                "description": f"Knowledge base collection: {collection_name}",
                "embedding_model": self.embedding_model,
                "dimensions": expected_dimensions
            }
        )
        
        self.logger.info(
            f"新しいコレクションを作成: {collection_name}",
            extra={
                "embedding_model": self.embedding_model,
                "dimensions": expected_dimensions
            }
        )
    
    def _read_pdf_file(self, file_path: Path) -> str:
        """
        PDFファイルからテキストを抽出
        
        Args:
            file_path: PDFファイルパス
            
        Returns:
            str: 抽出されたテキスト
            
        Raises:
            IndexingError: PDF読み込みエラー
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise IndexingError(
                        f"空のPDFファイル: {file_path}",
                        error_code="IDX-001",
                        details={"file_path": str(file_path)}
                    )
                
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                if not text_content.strip():
                    raise IndexingError(
                        f"PDFからテキストを抽出できませんでした: {file_path}",
                        error_code="IDX-001",
                        details={"file_path": str(file_path)}
                    )
                
                return text_content.strip()
                
        except IndexingError:
            raise
        except Exception as e:
            raise IndexingError(
                f"PDFファイル読み込みエラー: {file_path} - {e}",
                error_code="IDX-002",
                details={"file_path": str(file_path), "original_error": str(e)}
            ) from e
    
    def _read_txt_file(self, file_path: Path) -> str:
        """
        TXTファイルからテキストを読み込み
        
        Args:
            file_path: TXTファイルパス
            
        Returns:
            str: ファイル内容
            
        Raises:
            IndexingError: ファイル読み込みエラー
        """
        encodings = ['utf-8', 'shift_jis', 'euc-jp', 'cp932']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise IndexingError(
                    f"TXTファイル読み込みエラー: {file_path} - {e}",
                    error_code="IDX-003",
                    details={"file_path": str(file_path), "original_error": str(e)}
                ) from e
        
        raise IndexingError(
            f"文字エンコードエラー - サポートされていない文字コード: {file_path}",
            error_code="IDX-003",
            details={"file_path": str(file_path), "tried_encodings": encodings}
        )
    
    def _read_docx_file(self, file_path: Path) -> str:
        """
        DOCXファイルからテキストを読み込み
        
        Args:
            file_path: DOCXファイルパス
            
        Returns:
            str: ファイル内容
            
        Raises:
            IndexingError: ファイル読み込みエラー
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            full_text = []
            
            # 段落のテキストを抽出
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # テーブルのテキストも抽出
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return "\n".join(full_text)
            
        except ImportError:
            raise IndexingError(
                f"python-docxライブラリが必要です。pip install python-docx でインストールしてください",
                error_code="IDX-004",
                details={"file_path": str(file_path)}
            )
        except Exception as e:
            raise IndexingError(
                f"DOCXファイル読み込みエラー: {file_path} - {e}",
                error_code="IDX-004",
                details={"file_path": str(file_path), "original_error": str(e)}
            ) from e
    
    def _read_markdown_file(self, file_path: Path) -> str:
        """
        Markdownファイルからテキストを読み込み
        
        Args:
            file_path: Markdownファイルパス
            
        Returns:
            str: ファイル内容（Markdown記法をプレーンテキストに変換）
            
        Raises:
            IndexingError: ファイル読み込みエラー
        """
        try:
            from langchain_community.document_loaders import UnstructuredMarkdownLoader
            
            # UnstructuredMarkdownLoaderを使用してMarkdownを読み込み
            loader = UnstructuredMarkdownLoader(str(file_path))
            documents = loader.load()
            
            if not documents:
                raise IndexingError(
                    f"Markdownファイルにコンテンツが見つかりません: {file_path}",
                    error_code="IDX-009",
                    details={"file_path": str(file_path)}
                )
            
            # 複数のドキュメントが返される場合は結合
            full_text = '\n\n'.join([doc.page_content for doc in documents if doc.page_content.strip()])
            
            if not full_text.strip():
                raise IndexingError(
                    f"Markdownファイルから有効なテキストを抽出できませんでした: {file_path}",
                    error_code="IDX-009",
                    details={"file_path": str(file_path)}
                )
            
            self.logger.debug(f"Markdownファイル読み込み完了: {file_path.name}", extra={
                "file_path": str(file_path),
                "content_length": len(full_text),
                "documents_count": len(documents)
            })
            
            return full_text
            
        except Exception as e:
            # UnstructuredMarkdownLoaderが使用できない場合のフォールバック
            if "UnstructuredMarkdownLoader" in str(e) or "unstructured" in str(e).lower():
                return self._read_markdown_file_fallback(file_path)
            
            raise IndexingError(
                f"Markdownファイル読み込みエラー: {file_path} - {e}",
                error_code="IDX-009",
                details={"file_path": str(file_path), "original_error": str(e)}
            ) from e
    
    def _read_markdown_file_fallback(self, file_path: Path) -> str:
        """
        Markdownファイル読み込みのフォールバック処理
        
        Args:
            file_path: Markdownファイルパス
            
        Returns:
            str: ファイル内容
        """
        try:
            import markdown
            
            # ファイルを読み込み
            content = self._read_txt_file(file_path)  # エンコーディング処理を再利用
            
            # MarkdownをHTMLに変換してからプレーンテキストに
            md = markdown.Markdown(extensions=['extra', 'codehilite'])
            html = md.convert(content)
            
            # HTMLタグを除去してプレーンテキストに変換
            import re
            plain_text = re.sub('<[^<]+?>', '', html)
            plain_text = re.sub(r'\n\s*\n', '\n\n', plain_text)  # 連続する空行を整理
            
            return plain_text.strip()
            
        except Exception as e:
            # 最終フォールバック: プレーンテキストとして読み込み
            self.logger.warning(f"Markdownパース失敗、プレーンテキストとして処理: {file_path}")
            return self._read_txt_file(file_path)
    
    def _split_text_into_chunks(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[str]:
        """
        テキストをチャンクに分割
        
        Args:
            text: 分割対象テキスト
            chunk_size: チャンクサイズ
            chunk_overlap: チャンク間のオーバーラップ
            
        Returns:
            List[str]: テキストチャンクリスト
        """
        if chunk_size != 1000 or chunk_overlap != 200:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len
            )
            return text_splitter.split_text(text)
        
        return self.text_splitter.split_text(text)
    
    def _create_document_from_file(self, file_path: Path) -> Optional[Document]:
        """
        ファイルからDocumentオブジェクトを作成
        
        Args:
            file_path: ファイルパス
            
        Returns:
            Optional[Document]: 作成されたDocumentオブジェクト（失敗時はNone）
        """
        try:
            # ファイル存在確認
            if not file_path.exists():
                self.logger.warning(f"ファイルが存在しません: {file_path}")
                return None
            
            # ファイル拡張子を取得
            file_ext = file_path.suffix.lower()
            
            # ファイル内容を読み込み
            content = ""
            title = file_path.stem  # ファイル名（拡張子なし）をタイトルとする
            
            if file_ext == ".pdf":
                content = self._read_pdf_file(file_path)
            elif file_ext == ".txt":
                content = self._read_txt_file(file_path)
            elif file_ext == ".md":
                content = self._read_markdown_file(file_path)  # Markdownファイル専用読み込み
            elif file_ext == ".docx":
                content = self._read_docx_file(file_path)
            else:
                self.logger.warning(f"サポートされていないファイル形式: {file_ext}")
                return None
            
            if not content or not content.strip():
                self.logger.warning(f"ファイルの内容が空です: {file_path}")
                return None
            
            # Documentオブジェクトを作成
            document = Document.create_new(
                title=title,
                content=content,
                file_path=str(file_path),
                file_size=file_path.stat().st_size
            )
            
            self.logger.debug(f"ドキュメント作成成功: {file_path.name}")
            return document
            
        except Exception as e:
            self.logger.error(f"ドキュメント作成エラー: {file_path} - {e}", exc_info=True)
            return None
    
    def _create_embeddings(self, text_chunks: List[str]) -> List[List[float]]:
        """
        テキストチャンクから埋め込みベクトルを生成（最適化・エラーハンドリング強化）
        
        Args:
            text_chunks: テキストチャンクリスト
            
        Returns:
            List[List[float]]: 埋め込みベクトルリスト
        """
        if not text_chunks:
            return []
        
        if self._embedding_function is None:
            # デフォルト埋め込み (テスト・フォールバック用)
            self.logger.debug("デフォルト埋め込みを使用")
            return [[0.1 * (i + hash(chunk) % 100)] * 384 for i, chunk in enumerate(text_chunks)]
        
        try:
            # チャンク数が多い場合は分割処理
            if len(text_chunks) > 100:
                self.logger.info(f"大量チャンク処理: {len(text_chunks)}件を分割処理")
                batch_size = 50
                all_embeddings = []
                
                for i in range(0, len(text_chunks), batch_size):
                    batch = text_chunks[i:i + batch_size]
                    batch_embeddings = self._embedding_function.embed_documents(batch)
                    all_embeddings.extend(batch_embeddings)
                    
                    # キャンセルチェック
                    self.check_cancellation()
                
                return all_embeddings
            else:
                embeddings = self._embedding_function.embed_documents(text_chunks)
                return embeddings
                
        except Exception as e:
            self.logger.warning(
                f"埋め込み生成エラー、デフォルト埋め込みを使用: {e}",
                extra={"chunks_count": len(text_chunks), "error": str(e)}
            )
            # フォールバック: デフォルト埋め込み
            return [[0.1 * (i + hash(chunk) % 100)] * 384 for i, chunk in enumerate(text_chunks)]
    
    def add_document(self, document: Document) -> str:
        """
        ドキュメントをインデックスに追加
        
        Args:
            document: 追加するドキュメント
            
        Returns:
            str: 追加されたドキュメントID
            
        Raises:
            IndexingError: インデックス追加エラー
        """
        try:
            self.check_cancellation()
            
            # ISSUE-027対応: 埋め込み次元数の互換性チェック
            compatibility_result = self.check_embedding_dimension_compatibility()
            if not compatibility_result['is_compatible']:
                if compatibility_result.get('needs_recreation', False):
                    current_dim = compatibility_result.get('current_dimensions', '不明')
                    expected_dim = compatibility_result.get('expected_dimensions', '不明')
                    self.logger.warning(
                        f"次元数不整合を検出、コレクションを再作成します: "
                        f"現在={current_dim}, 期待={expected_dim}",
                        extra=compatibility_result
                    )
                    self._recreate_collection_with_new_dimensions()
                else:
                    error_msg = compatibility_result.get('error', '次元数互換性チェックに失敗')
                    raise IndexingError(
                        f"埋め込み次元数互換性エラー: {error_msg}",
                        error_code="IDX-027",
                        details=compatibility_result
                    )
            
            # テキスト内容を取得
            if not document.content:
                if document.file_path:
                    file_path = Path(document.file_path)
                    
                    # セキュリティ検証
                    is_valid, error_msg = self.file_validator.validate_file(str(file_path))
                    if not is_valid:
                        raise IndexingError(
                            f"ファイル検証エラー: {error_msg}",
                            error_code="IDX-008",
                            details={"file_path": str(file_path), "validation_error": error_msg}
                        )
                    
                    if document.file_type.lower() == 'pdf':
                        content = self._read_pdf_file(file_path)
                    elif document.file_type.lower() == 'txt':
                        content = self._read_txt_file(file_path)
                    elif document.file_type.lower() in ['md', 'markdown']:
                        content = self._read_markdown_file(file_path)
                    elif document.file_type.lower() == 'docx':
                        content = self._read_docx_file(file_path)
                    else:
                        raise IndexingError(
                            f"サポートされていないファイル形式: {document.file_type}",
                            error_code="IDX-005",
                            details={"file_type": document.file_type, "file_path": document.file_path}
                        )
                    # ドキュメントのcontentを更新
                    document.content = content
                else:
                    raise IndexingError(
                        "ドキュメントの内容またはファイルパスが必要です",
                        error_code="IDX-006"
                    )
            
            # テキストをチャンクに分割
            text_chunks = self._split_text_into_chunks(document.content)
            if not text_chunks:
                raise IndexingError(
                    f"テキストチャンクの生成に失敗: {document.file_path}",
                    error_code="IDX-007",
                    details={"document_filename": Path(document.file_path).name}
                )
            
            self.check_cancellation()
            
            # 埋め込みベクトルを生成
            embeddings = self._create_embeddings(text_chunks)
            
            self.check_cancellation()
            
            # ドキュメントIDを生成
            document_id = document.id
            
            # メタデータを準備
            metadatas = []
            ids = []
            for i, chunk in enumerate(text_chunks):
                chunk_id = f"{document_id}_{i}"
                metadata = {
                    "document_filename": Path(document.file_path).name,
                    "file_path": document.file_path or "",
                    "file_type": document.file_type,
                    "file_size": document.file_size,
                    "chunk_index": i,
                    "document_id": document_id,
                    "created_at": document.created_at.isoformat() if document.created_at else ""
                }
                metadatas.append(metadata)
                ids.append(chunk_id)
            
            # ChromaDBコレクションに追加
            self.collection.add(
                documents=text_chunks,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            self.logger.info(f"ドキュメントインデックス追加完了: {Path(document.file_path).name}", extra={
                "document_id": document_id,
                "document_filename": Path(document.file_path).name,
                "chunks_count": len(text_chunks)
            })
            
            return document_id
            
        except IndexingError:
            raise
        except Exception as e:
            # ChromaDB次元数不一致エラーの検出
            error_message = str(e)
            if "Collection expecting embedding with dimension" in error_message:
                # 次元数不一致エラーメッセージを抽出
                if "got" in error_message and "expecting" in error_message:
                    # エラーメッセージから既存の次元数を推定
                    try:
                        # "expecting embedding with dimension of X, got Y" パターンを解析
                        parts = error_message.split("expecting embedding with dimension of ")[1].split(", got ")
                        expected_dim = int(parts[0])
                        current_dim = int(parts[1])
                        
                        # 次元数から推定される埋め込みモデル名を取得
                        dimension_model_map = {
                            768: "nomic-embed-text",
                            1024: "mxbai-embed-large"
                        }
                        
                        suggested_model = dimension_model_map.get(expected_dim, "不明")
                        
                        user_friendly_message = f"異なる埋め込みモデルでコレクションが作成されています。埋め込みモデル「{suggested_model}」に変更してインデックスを作成してください。"
                        
                        raise IndexingError(
                            user_friendly_message,
                            error_code="IDX-009",
                            details={
                                "document_filename": Path(document.file_path).name, 
                                "expected_dimensions": expected_dim,
                                "current_dimensions": current_dim,
                                "suggested_model": suggested_model,
                                "original_error": error_message
                            }
                        ) from e
                    except (ValueError, IndexError):
                        # パース失敗時は汎用メッセージ
                        raise IndexingError(
                            "異なる埋め込みモデルでコレクションが作成されています。設定で埋め込みモデルを確認し、適切なモデルに変更してインデックスを作成してください。",
                            error_code="IDX-009",
                            details={"document_filename": Path(document.file_path).name, "original_error": error_message}
                        ) from e
            
            # その他のエラー
            raise IndexingError(
                f"ドキュメントインデックス追加エラー: {Path(document.file_path).name} - {e}",
                error_code="IDX-008",
                details={"document_filename": Path(document.file_path).name, "original_error": str(e)}
            ) from e
    
    def add_documents(self, documents: List[Document]) -> List[str]:
        """
        複数のドキュメントを一括でインデックスに追加
        
        Args:
            documents: 追加するドキュメントリスト
            
        Returns:
            List[str]: 追加されたドキュメントIDリスト
        """
        document_ids = []
        total_docs = len(documents)
        
        # 進捗表示が必要な場合のみProgressTrackerを使用
        progress_tracker = None
        if should_show_progress(total_docs * 2):  # 1ドキュメントあたり約2秒と仮定
            progress_tracker = ProgressTracker(
                total=total_docs,
                description="ドキュメントインデックス作成"
            )
        
        try:
            for i, document in enumerate(documents):
                self.check_cancellation()
                
                if progress_tracker:
                    progress_tracker.update(
                        message=f"処理中: {Path(document.file_path).name} ({i+1}/{total_docs})"
                    )
                
                document_id = self.add_document(document)
                document_ids.append(document_id)
            
            if progress_tracker:
                progress_tracker.finish("インデックス作成完了")
            
            return document_ids
            
        except Exception as e:
            if progress_tracker:
                progress_tracker.cancel()
            raise
    
    def search_documents(
        self,
        query: str,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        ドキュメントを検索
        
        Args:
            query: 検索クエリ
            top_k: 上位K件を取得
            
        Returns:
            List[Dict[str, Any]]: 検索結果リスト
        """
        try:
            # クエリの埋め込みベクトルを生成
            query_embedding = self._create_embeddings([query])[0]
            
            # ChromaDBで検索実行
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )
            
            # 結果を整形
            search_results = []
            for i in range(len(results['documents'][0])):
                search_results.append({
                    'content': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i]
                })
            
            self.logger.info(f"ドキュメント検索完了", extra={
                "query": query,
                "results_count": len(search_results),
                "top_k": top_k
            })
            
            return search_results
            
        except Exception as e:
            raise IndexingError(
                f"ドキュメント検索エラー: {e}",
                error_code="IDX-009",
                details={"query": query, "top_k": top_k, "original_error": str(e)}
            ) from e
    
    def update_document(self, document_id: str, updated_document: Document) -> bool:
        """
        ドキュメントを更新
        
        Args:
            document_id: 更新対象ドキュメントID
            updated_document: 更新後のドキュメント
            
        Returns:
            bool: 更新成功フラグ
        """
        try:
            # 既存のドキュメントを削除
            self.delete_document(document_id)
            
            # 新しいドキュメントを追加 (IDを維持)
            updated_document.id = document_id
            self.add_document(updated_document)
            
            return True
            
        except Exception as e:
            raise IndexingError(
                f"ドキュメント更新エラー: {document_id} - {e}",
                error_code="IDX-010",
                details={"document_id": document_id, "original_error": str(e)}
            ) from e
    
    def delete_document(self, document_id: str) -> bool:
        """
        ドキュメントをインデックスから削除
        
        Args:
            document_id: 削除対象ドキュメントID
            
        Returns:
            bool: 削除成功フラグ
        """
        try:
            # ドキュメントのチャンクIDを検索
            results = self.collection.get(
                where={"document_id": document_id},
                include=["metadatas"]
            )
            
            if not results['ids']:
                raise IndexingError(
                    f"削除対象ドキュメントが見つかりません: {document_id}",
                    error_code="IDX-004",
                    details={"document_id": document_id}
                )
            
            # 関連するチャンクIDをすべて削除
            chunk_ids = results['ids']
            self.collection.delete(ids=chunk_ids)
            
            self.logger.info(f"ドキュメント削除完了: {document_id}", extra={
                "document_id": document_id,
                "deleted_chunks": len(chunk_ids)
            })
            
            return True
            
        except IndexingError:
            raise
        except Exception as e:
            raise IndexingError(
                f"ドキュメント削除エラー: {document_id} - {e}",
                error_code="IDX-004",
                details={"document_id": document_id, "original_error": str(e)}
            ) from e
    
    def get_model_expected_dimensions(self) -> int:
        """
        埋め込みモデルの期待次元数を取得
        
        Returns:
            int: 期待される次元数
        """
        expected_dims = self.EMBEDDING_DIMENSIONS.get(self.embedding_model)
        if expected_dims is None:
            # 未知のモデルの場合は実際にテストして取得
            try:
                test_embedding = self._embedding_function.embed_query("test")
                expected_dims = len(test_embedding)
                self.logger.warning(
                    f"未知の埋め込みモデル、実際の次元数を使用: {self.embedding_model} -> {expected_dims}"
                )
            except Exception as e:
                self.logger.error(f"次元数取得失敗: {e}")
                expected_dims = 384  # デフォルト値
        
        return expected_dims

    def update_embedding_model(self, new_embedding_model: str) -> bool:
        """
        埋め込みモデルを変更し、必要に応じてコレクションを再作成
        
        Args:
            new_embedding_model: 新しい埋め込みモデル名
            
        Returns:
            bool: 更新成功フラグ
        """
        try:
            old_model = self.embedding_model
            old_dimensions = self.get_model_expected_dimensions()
            
            # モデルを変更
            self.embedding_model = new_embedding_model
            
            # 新しいモデルの埋め込み関数を初期化
            self._initialize_embedding_function()
            
            # 新しいモデルの次元数を取得
            new_dimensions = self.get_model_expected_dimensions()
            
            # 次元数が異なる場合はコレクションを再作成
            if old_dimensions != new_dimensions:
                self.logger.info(
                    f"埋め込みモデル変更による次元数変更を検出: {old_model}({old_dimensions}) -> {new_embedding_model}({new_dimensions})"
                )
                
                # 既存のドキュメント数を確認
                doc_count = self.collection.count() if self.collection else 0
                
                if doc_count > 0:
                    self.logger.warning(
                        f"既存の{doc_count}件のドキュメントが削除されます。コレクションを再作成中..."
                    )
                
                # コレクションを再作成
                self._recreate_collection_with_new_dimensions()
                
                return True
            else:
                self.logger.info(f"埋め込みモデルを更新: {old_model} -> {new_embedding_model}（次元数変更なし: {new_dimensions}）")
                return False
                
        except Exception as e:
            self.logger.error(f"埋め込みモデル更新エラー: {e}")
            # エラー時は元のモデルに戻す
            self.embedding_model = old_model if 'old_model' in locals() else "nomic-embed-text"
            self._initialize_embedding_function()
            raise IndexingError(
                f"埋め込みモデル更新エラー: {e}",
                error_code="IDX-MODEL-UPDATE",
                details={"old_model": old_model if 'old_model' in locals() else None, 
                        "new_model": new_embedding_model, "original_error": str(e)}
            ) from e

    def check_embedding_dimension_compatibility(self) -> dict:
        """
        埋め込み次元数の互換性をチェック (ISSUE-027対応)
        
        Returns:
            dict: 互換性チェック結果
        """
        try:
            result = {
                'is_compatible': True,
                'current_dimensions': None,
                'expected_dimensions': None,
                'needs_recreation': False,
                'error': None,
                'model_name': self.embedding_model
            }
            
            # 現在の埋め込みモデルから次元数を取得
            try:
                if self._embedding_function is None:
                    result['current_dimensions'] = self.get_model_expected_dimensions()
                else:
                    test_embedding = self._embedding_function.embed_query("test")
                    result['current_dimensions'] = len(test_embedding)
            except Exception as e:
                result['error'] = f"埋め込みモデルテストに失敗: {e}"
                result['is_compatible'] = False
                return result
            
            # ChromaDBコレクションから期待される次元数を取得
            count = self.collection.count()
            if count > 0:
                sample = self.collection.get(limit=1, include=['embeddings'])
                embeddings = sample.get('embeddings', [])
                if embeddings is not None and len(embeddings) > 0:
                    first_embedding = embeddings[0]
                    if first_embedding is not None and len(first_embedding) > 0:
                        result['expected_dimensions'] = len(first_embedding)
                        
                        # 次元数の比較
                        if result['current_dimensions'] != result['expected_dimensions']:
                            result['is_compatible'] = False
                            result['needs_recreation'] = True
                            result['error'] = f"次元数不整合: 現在{result['current_dimensions']}次元({self.embedding_model}) vs コレクション{result['expected_dimensions']}次元"
            else:
                # コレクションが空の場合は互換性チェック不要
                result['expected_dimensions'] = result['current_dimensions']
                
            self.logger.info(f"埋め込み次元数互換性チェック完了", extra=result)
            return result
            
        except Exception as e:
            error_result = {
                'is_compatible': False,
                'error': f"互換性チェックエラー: {e}",
                'needs_recreation': True,
                'model_name': self.embedding_model
            }
            self.logger.error("埋め込み次元数互換性チェックに失敗", extra=error_result)
            return error_result
    
    def recreate_collection_if_incompatible(self) -> bool:
        """
        次元数不整合時にコレクションを再作成 (ISSUE-027対応)
        
        Returns:
            bool: 再作成実行フラグ
        """
        try:
            compatibility = self.check_embedding_dimension_compatibility()
            
            if compatibility['needs_recreation']:
                self.logger.warning(
                    f"次元数不整合を検出、コレクションを再作成します: {compatibility['error']}"
                )
                
                # コレクション削除
                self.client.delete_collection(self.collection_name)
                
                # 新しいコレクションを作成
                self.collection = self.client.create_collection(
                    name=self.collection_name,
                    embedding_function=self._embedding_function,
                    metadata={"hnsw:space": "cosine"}
                )
                
                self.logger.info(f"コレクション再作成完了: 新しい次元数 {compatibility['current_dimensions']}")
                return True
            else:
                self.logger.info("次元数互換性は正常です")
                return False
                
        except Exception as e:
            raise IndexingError(
                f"コレクション再作成エラー: {e}",
                error_code="IDX_DIMENSION_RECREATE",
                details={"original_error": str(e)}
            ) from e
    
    def _recreate_collection_with_new_dimensions(self) -> None:
        """
        次元数不整合時にコレクションを再作成 (ISSUE-027対応)
        
        Raises:
            IndexingError: コレクション再作成エラー
        """
        try:
            # 現在のコレクション情報を保存
            old_doc_count = self.collection.count() if self.collection else 0
            
            self.logger.info(
                f"コレクション再作成開始: {self.collection_name}",
                extra={"old_document_count": old_doc_count}
            )
            
            # コレクション削除
            try:
                self.client.delete_collection(self.collection_name)
                self.logger.info(f"旧コレクション削除完了: {self.collection_name}")
            except Exception as e:
                self.logger.warning(f"旧コレクション削除時の警告: {e}")
            
            # 新しいコレクションを作成
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            # 新しい次元数を取得して確認
            test_embedding = self._embedding_function.embed_query("test")
            new_dimensions = len(test_embedding)
            
            self.logger.info(
                f"コレクション再作成完了: {self.collection_name}",
                extra={
                    "new_dimensions": new_dimensions,
                    "old_document_count": old_doc_count,
                    "collection_name": self.collection_name
                }
            )
            
        except Exception as e:
            raise IndexingError(
                f"コレクション再作成エラー: {e}",
                error_code="IDX-027-RECREATE",
                details={"collection_name": self.collection_name, "original_error": str(e)}
            ) from e
    
    def clear_collection(self) -> bool:
        """
        コレクション内の全ドキュメントをクリア
        
        ChromaDB v1.0.17以降では空のwhere条件による削除が制限されているため、
        コレクション再作成による全削除を実行
        
        Returns:
            bool: クリア成功フラグ
        """
        try:
            # コレクション内のドキュメント数を確認
            doc_count = self.collection.count()
            
            if doc_count == 0:
                self.logger.info(f"コレクションは既に空です", extra={
                    "collection_name": self.collection_name
                })
                return True
            
            # ChromaDB v1.0.17以降での全削除方法
            # 方法1: 全ドキュメントのIDを取得して削除
            try:
                # 全IDを取得
                result = self.collection.get()
                all_ids = result.get('ids', [])
                
                if all_ids:
                    # IDによる削除実行
                    self.collection.delete(ids=all_ids)
                    self.logger.info(f"IDによる全削除完了", extra={
                        "collection_name": self.collection_name,
                        "deleted_documents": len(all_ids)
                    })
                else:
                    self.logger.info(f"削除対象のIDが見つかりませんでした", extra={
                        "collection_name": self.collection_name
                    })
                    
            except Exception as delete_error:
                # 方法2: コレクション再作成による全削除
                self.logger.warning(f"ID削除失敗、コレクション再作成を実行: {delete_error}")
                
                # 現在のコレクション設定を保存
                original_name = self.collection_name
                
                # コレクションを削除
                self.client.delete_collection(name=original_name)
                
                # 同じ名前でコレクションを再作成
                self.collection = self.client.create_collection(
                    name=original_name,
                    embedding_function=self._embedding_function,
                    metadata={"hnsw:space": "cosine"}
                )
                
                self.logger.info(f"コレクション再作成による全削除完了", extra={
                    "collection_name": self.collection_name,
                    "deleted_documents": doc_count
                })
            
            return True
            
        except Exception as e:
            raise IndexingError(
                f"コレクションクリアエラー: {e}",
                error_code="IDX-011",
                details={"collection_name": self.collection_name, "original_error": str(e)}
            ) from e
    
    def rebuild_from_directory(self, directory_path: Path) -> List[str]:
        """
        ディレクトリからインデックスを再構築
        
        Args:
            directory_path: 処理対象ディレクトリ
            
        Returns:
            List[str]: 追加されたドキュメントIDリスト
        """
        try:
            # サポート対象ファイルを収集
            supported_extensions = ['*.pdf', '*.txt']
            file_paths = []
            
            for ext in supported_extensions:
                file_paths.extend(directory_path.glob(f"**/{ext}"))
            
            if not file_paths:
                self.logger.warning(f"処理対象ファイルが見つかりません: {directory_path}")
                return []
            
            # ドキュメントオブジェクトを作成
            documents = []
            for file_path in file_paths:
                doc = self._create_document_from_file(file_path)
                if doc:  # ファイル読み込みが成功した場合のみ追加
                    documents.append(doc)
            
            # 一括でインデックスに追加
            document_ids = self.add_documents(documents)
            
            self.logger.info(f"ディレクトリからインデックス再構築完了", extra={
                "directory_path": str(directory_path),
                "processed_files": len(file_paths),
                "document_ids": len(document_ids)
            })
            
            return document_ids
            
        except Exception as e:
            raise IndexingError(
                f"ディレクトリインデックス再構築エラー: {directory_path} - {e}",
                error_code="IDX-012",
                details={"directory_path": str(directory_path), "original_error": str(e)}
            ) from e
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """
        コレクション統計情報を取得
        
        Returns:
            Dict[str, Any]: コレクション統計情報
        """
        try:
            doc_count = self.collection.count()
            
            return {
                "collection_name": self.collection_name,
                "document_count": doc_count,
                "db_path": str(self.db_path),
                "embedding_model": self.embedding_model
            }
            
        except Exception as e:
            self.logger.warning(f"統計情報取得エラー: {e}")
            return {
                "collection_name": self.collection_name,
                "document_count": 0,
                "db_path": str(self.db_path),
                "embedding_model": self.embedding_model,
                "error": str(e)
            }
    
    def rebuild_index_from_folders(self, folder_paths: List[str]) -> bool:
        """
        指定されたフォルダパスから文書を読み込み、インデックスを再構築します。

        Args:
            folder_paths: インデックス化するフォルダのパスリスト

        Returns:
            bool: 成功した場合True

        Raises:
            IndexingError: インデックス再構築に失敗した場合
        """
        self.logger.info(f"フォルダからインデックス再構築開始: {folder_paths}")
        try:
            self.clear_collection() # 既存のインデックスをクリア

            total_files_to_process = 0
            for folder_path_str in folder_paths:
                folder_path = Path(folder_path_str)
                if folder_path.is_dir():
                    # サポート対象ファイルを収集 (rebuild_from_directoryと同じロジック)
                    supported_extensions = ['*.pdf', '*.txt']
                    for ext in supported_extensions:
                        total_files_to_process += len(list(folder_path.glob(f"**/{ext}")))
                else:
                    self.logger.warning(f"指定されたパスはディレクトリではありません: {folder_path_str}")

            if total_files_to_process == 0:
                self.logger.info("処理対象ファイルが見つかりませんでした。")
                return True

            progress_tracker = ProgressTracker(
                total=total_files_to_process,
                description="フォルダからインデックスを再構築中"
            )

            processed_files_count = 0
            for folder_path_str in folder_paths:
                folder_path = Path(folder_path_str)
                if folder_path.is_dir():
                    self.check_cancellation()
                    self.logger.info(f"フォルダ処理中: {folder_path_str}")
                    
                    # rebuild_from_directoryのロジックをここに統合
                    supported_extensions = ['*.pdf', '*.txt']
                    file_paths = []
                    for ext in supported_extensions:
                        file_paths.extend(folder_path.glob(f"**/{ext}"))

                    for file_path in file_paths:
                        self.check_cancellation()
                        doc = self._create_document_from_file(file_path)
                        if not doc:  # ファイル読み込みに失敗した場合はスキップ
                            continue
                        self.add_document(doc) # 個別にドキュメントを追加
                        processed_files_count += 1
                        progress_tracker.update(
                            processed_files_count,
                            message=f"処理中: {file_path.name} ({processed_files_count}/{total_files_to_process})"
                        )
                
            progress_tracker.finish("インデックス再構築完了")
            self.logger.info("フォルダからのインデックス再構築完了")
            return True

        except Exception as e:
            self.logger.error(f"フォルダからのインデックス再構築エラー: {e}")
            if 'progress_tracker' in locals() and progress_tracker:
                progress_tracker.cancel()
            raise IndexingError(f"フォルダからのインデックス再構築に失敗しました: {e}")